import streamlit as st
import anthropic
import pdfplumber
import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(page_title="PDF Translator", page_icon="📄", layout="centered")

st.title("📄 PDF → Word Translator")
st.caption("Extracts paragraphs from a PDF, translates them, and exports a 2-column Word document.")

# ── Sidebar: settings ──────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Settings")
    api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
    st.markdown("---")
    source_lang = st.selectbox("Source language", [
        "Auto-detect", "English", "Romanian", "French", "German", "Spanish",
        "Italian", "Portuguese", "Dutch", "Polish", "Hungarian", "Czech",
        "Russian", "Ukrainian", "Turkish", "Arabic", "Chinese", "Japanese", "Korean"
    ])
    target_lang = st.selectbox("Target language", [
        "English", "Romanian", "French", "German", "Spanish",
        "Italian", "Portuguese", "Dutch", "Polish", "Hungarian", "Czech",
        "Russian", "Ukrainian", "Turkish", "Arabic", "Chinese", "Japanese", "Korean"
    ])
    st.markdown("---")
    st.markdown("**Left column** → Original text  \n**Right column** → Translation")

# ── Helpers ────────────────────────────────────────────────────────────────────

def extract_paragraphs(pdf_file) -> list[str]:
    """Extract non-empty paragraphs from a PDF."""
    paragraphs = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for para in text.split("\n\n"):
                clean = para.strip().replace("\n", " ")
                clean = re.sub(r"\s{2,}", " ", clean)
                if len(clean) > 20:          # skip very short fragments
                    paragraphs.append(clean)
    return paragraphs


def translate_paragraphs(client, paragraphs: list[str], source: str, target: str) -> list[str]:
    """Translate a list of paragraphs via Claude, preserving order."""
    numbered = "\n\n".join(f"[{i+1}] {p}" for i, p in enumerate(paragraphs))
    from_clause = "" if source == "Auto-detect" else f" from {source}"
    prompt = (
        f"Translate the following numbered paragraphs{from_clause} to {target}. "
        "Return ONLY the translations, each prefixed with the same [number] tag, "
        "in the same order. Do not add any explanation or preamble.\n\n"
        f"{numbered}"
    )
    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.content[0].text.strip()
    # Parse [N] tags back into list
    parts = re.split(r"\[\d+\]\s*", raw)
    translations = [p.strip() for p in parts if p.strip()]
    # Pad if mismatch
    while len(translations) < len(paragraphs):
        translations.append("")
    return translations[:len(paragraphs)]


def set_cell_background(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_docx(paragraphs: list[str], translations: list[str],
               source: str, target: str) -> bytes:
    """Build a .docx with a 2-column table."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Title
    title = doc.add_heading("PDF Translation", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub_text = f"Source: {source}   →   Target: {target}"
    sub = doc.add_paragraph(sub_text)
    sub.runs[0].font.size   = Pt(10)
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths (equal split)
    for cell in table.columns[0].cells:
        cell.width = Inches(3.2)
    for cell in table.columns[1].cells:
        cell.width = Inches(3.2)

    # Header row
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, [source if source != "Auto-detect" else "Original", target]):
        cell.text = label
        set_cell_background(cell, "1F3864")
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(11)
        run.font.name  = "Arial"

    # Data rows
    for i, (orig, trans) in enumerate(zip(paragraphs, translations)):
        row = table.add_row().cells
        bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        for cell, content in zip(row, [orig, trans]):
            cell.text = content
            set_cell_background(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.name = "Arial"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Main UI ────────────────────────────────────────────────────────────────────

uploaded = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded and not api_key:
    st.warning("Please enter your Anthropic API key in the sidebar.")

if uploaded and api_key:
    st.success(f"File loaded: **{uploaded.name}**")

    if st.button("🚀 Extract, Translate & Export", type="primary", use_container_width=True):
        with st.spinner("Extracting paragraphs…"):
            try:
                paragraphs = extract_paragraphs(uploaded)
            except Exception as e:
                st.error(f"PDF extraction failed: {e}")
                st.stop()

        if not paragraphs:
            st.error("No text could be extracted from this PDF. It may be a scanned/image-only PDF.")
            st.stop()

        st.info(f"Found **{len(paragraphs)}** paragraphs. Translating…")

        with st.spinner("Translating via Claude…"):
            try:
                client = anthropic.Anthropic(api_key=api_key)
                # Translate in batches of 20 to stay within token limits
                batch_size = 20
                all_translations = []
                progress = st.progress(0)
                for start in range(0, len(paragraphs), batch_size):
                    batch = paragraphs[start:start + batch_size]
                    translations = translate_paragraphs(client, batch, source_lang, target_lang)
                    all_translations.extend(translations)
                    progress.progress(min((start + batch_size) / len(paragraphs), 1.0))
            except Exception as e:
                st.error(f"Translation failed: {e}")
                st.stop()

        with st.spinner("Building Word document…"):
            docx_bytes = build_docx(paragraphs, all_translations, source_lang, target_lang)

        st.success("✅ Done! Your document is ready.")

        st.download_button(
            label="📥 Download Word Document",
            data=docx_bytes,
            file_name=f"{uploaded.name.replace('.pdf','')}_translated.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        # Preview table
        with st.expander("👁 Preview (first 5 paragraphs)"):
            for orig, trans in zip(paragraphs[:5], all_translations[:5]):
                col1, col2 = st.columns(2)
                col1.write(orig)
                col2.write(trans)
                st.divider()
